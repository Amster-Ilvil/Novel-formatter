#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build a vertical DOCX with column provenance verification."""
from __future__ import annotations

import argparse
import os
import re
import sys
from collections import Counter
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from models.document import UnifiedDocument, Block, BlockType


def _set_vertical_layout(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    sect_pr = doc.sections[0]._sectPr
    text_dir = OxmlElement('w:textDirection')
    text_dir.set(qn('w:val'), 'tbRl')
    sect_pr.append(text_dir)


def _add_page_break(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _block_text_for_word(b: Block) -> str:
    if b.type == BlockType.RUBY:
        return re.sub(r'([^\s|]+)\|([^\s|]+)', r'\1（\2）', b.text)
    return b.text


def _source_column_ids(block: Block) -> list[str]:
    metadata = block.metadata or {}
    values = metadata.get("source_column_ids") or []
    if isinstance(values, str):
        values = [values]
    result = [str(value) for value in values if str(value)]
    column_id = str(metadata.get("column_id", ""))
    if column_id and column_id not in result:
        result.append(column_id)
    return result


def _expected_column_manifest(doc: UnifiedDocument) -> tuple[list[str], dict[str, int]]:
    audit = getattr(doc.metadata, "column_ocr_audit", {}) or {}
    pages = audit.get("pages", {}) if isinstance(audit, dict) else {}
    expected: list[str] = []
    page_by_id: dict[str, int] = {}
    for page_key, report in sorted(pages.items(), key=lambda item: int(item[0])):
        page_no = int(page_key)
        for column_id in report.get("column_ids", []) or []:
            column_id = str(column_id)
            if not column_id:
                continue
            expected.append(column_id)
            page_by_id[column_id] = page_no
    return expected, page_by_id


def _verify_docx_text(temp_path: Path, expected_texts: list[str]) -> None:
    from docx import Document
    reopened = Document(str(temp_path))
    actual = [paragraph.text for paragraph in reopened.paragraphs if paragraph.text]
    if actual != expected_texts:
        first = next(
            (index for index, (left, right) in enumerate(zip(expected_texts, actual), start=1) if left != right),
            min(len(expected_texts), len(actual)) + 1,
        )
        raise RuntimeError(
            "DOCX 保存后复读校验失败：写入前后的段落序列不一致，"
            f"首个差异位于第 {first} 段（预计 {len(expected_texts)}，复读 {len(actual)}）。"
        )


def build_word(
    doc: UnifiedDocument,
    output_path: str,
    vertical: bool = True,
    page_breaks: bool = True,
    verbose: bool = True,
) -> dict:
    """Generate DOCX atomically and verify every fixed-region source column."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    word = Document()
    if vertical:
        _set_vertical_layout(word)
    for paragraph in list(word.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)

    text_types = {
        BlockType.PARAGRAPH, BlockType.DIALOGUE, BlockType.CHAPTER,
        BlockType.SECTION, BlockType.RUBY,
    }
    expected_ids, page_by_id = _expected_column_manifest(doc)
    represented_ids: list[str] = []
    written_texts: list[str] = []
    prev_page: int | None = None

    for block in doc.blocks:
        if (block.metadata or {}).get("consumed"):
            continue
        if block.type not in text_types:
            continue
        text = _block_text_for_word(block)
        if not text.strip():
            continue
        if page_breaks and prev_page is not None and block.page != prev_page and written_texts:
            _add_page_break(word)
        prev_page = block.page
        paragraph = word.add_paragraph()
        run = paragraph.add_run(text)
        if block.type == BlockType.CHAPTER:
            run.bold = True
            run.font.size = Pt(16)
        written_texts.append(text)
        represented_ids.extend(_source_column_ids(block))

    if expected_ids:
        expected_set = set(expected_ids)
        represented_set = set(represented_ids)
        missing = [column_id for column_id in expected_ids if column_id not in represented_set]
        unexpected = [column_id for column_id in represented_set if column_id not in expected_set]
        audit = getattr(doc.metadata, "column_ocr_audit", {}) or {}
        model_ok = bool(audit.get("model_integrity_passed"))
        if not model_ok or missing or unexpected:
            details = []
            if not model_ok:
                details.append("OCR 文档模型阶段的列对账尚未通过")
            if missing:
                details.append("缺少：" + ", ".join(missing[:12]))
            if unexpected:
                details.append("额外：" + ", ".join(unexpected[:12]))
            raise RuntimeError(
                "DOCX 导出前列 ID 对账失败，已停止保存以避免末列丢失。\n"
                + "\n".join(f"• {item}" for item in details)
            )

    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    temp_path = target.with_name(f".{target.stem}.writing-{os.getpid()}.docx")
    try:
        word.save(str(temp_path))
        _verify_docx_text(temp_path, written_texts)
        temp_path.replace(target)
    finally:
        temp_path.unlink(missing_ok=True)

    # A Formatter may split one physical source column into multiple output
    # paragraphs.  DOCX coverage therefore counts unique source-column IDs,
    # while exact missing/extra ID checks still prevent a duplicated middle
    # column from concealing an omitted final column.
    represented_unique_ids = list(dict.fromkeys(represented_ids))
    per_page_docx: dict[int, int] = Counter(
        page_by_id[column_id] for column_id in represented_unique_ids if column_id in page_by_id
    )
    if expected_ids:
        audit = doc.metadata.column_ocr_audit
        for page_key, report in audit.get("pages", {}).items():
            page_no = int(page_key)
            report["docx_written"] = int(per_page_docx.get(page_no, 0))
            report["docx_passed"] = (
                int(report.get("expected", 0))
                == int(report.get("recognized", 0))
                == int(report.get("model_written", 0))
                == int(report.get("docx_written", 0))
            )
        totals = audit.setdefault("totals", {})
        totals["docx_written"] = len(represented_unique_ids)
        last_page = int(audit.get("last_ocr_page", 0) or 0)
        last_report = audit.get("pages", {}).get(str(last_page), {})
        audit["last_page_all_columns_exported"] = bool(last_report.get("docx_passed"))
        audit["docx_integrity_passed"] = (
            len(represented_unique_ids) == len(expected_ids)
            and all(bool(report.get("docx_passed")) for report in audit.get("pages", {}).values())
            and bool(audit["last_page_all_columns_exported"])
        )
        doc.metadata.column_ocr_integrity_passed = bool(audit["docx_integrity_passed"])

    report = {
        "path": str(target),
        "paragraphs_written": len(written_texts),
        "expected_columns": len(expected_ids),
        "docx_written_columns": len(represented_unique_ids),
        "source_column_references": len(represented_ids),
        "last_page_all_columns_exported": bool(
            (getattr(doc.metadata, "column_ocr_audit", {}) or {}).get("last_page_all_columns_exported", True)
        ),
    }
    if verbose:
        size_kb = target.stat().st_size // 1024
        audit_text = ""
        if expected_ids:
            audit_text = f", 列对账 {len(expected_ids)}/{len(represented_unique_ids)}"
        print(f"✅  Word 已生成: {target}  ({size_kb} KB, {len(written_texts)} 段{audit_text})")
    return report


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="UnifiedDocument → Word (.docx)")
    parser.add_argument("input_json", help="Formatter 输出的 JSON")
    parser.add_argument("output_docx", help="输出 .docx 路径")
    parser.add_argument("--horizontal", action="store_true", help="横排模式（默认竖排）")
    parser.add_argument("--no-page-breaks", action="store_true", help="不按原书页插入分页符")
    parser.add_argument("--quiet", "-q", action="store_true")
    args = parser.parse_args()
    with open(args.input_json, encoding="utf-8") as fh:
        document = UnifiedDocument.from_json(fh.read())
    build_word(
        document,
        output_path=args.output_docx,
        vertical=not args.horizontal,
        page_breaks=not args.no_page_breaks,
        verbose=not args.quiet,
    )
