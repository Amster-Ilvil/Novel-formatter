#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 导入适配器
导入 Abbyy FineReader / Adobe Acrobat OCR 输出的 DOCX 文件，
转换为 UnifiedDocument。

依赖：
    pip install python-docx

用法：
    python docx_adapter.py input.docx output.json
"""

from __future__ import annotations

import re
import sys
import os
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from models.document import UnifiedDocument, Block, BlockType, Metadata

CHAPTER_RE = re.compile(
    # フロローグ：常见的"プロローグ"半浊点丢失错读，一并纳入识别。
    r'^(序章|終章|プロローグ|フロローグ|エピローグ|後記|あとがき|'
    r'幕間[\s　]|幕間$|'
    r'第[一二三四五六七八九十百〇零\d]+[章話節巻]'
    r'|Chapter\s*\d+)',
    re.IGNORECASE
)

SECTION_RE = re.compile(
    r'^[◆※☆★●○＊◇■□▼▽△▲]{1,5}$'
)

DIALOGUE_START = ('「', '『', '（')
DIALOGUE_END = ('」', '』', '）')


def extract_images_from_docx(docx_path: str, output_dir: str) -> dict[str, str]:
    """
    从 DOCX 中提取嵌入图片，保存到 output_dir。
    返回 {rId: 图片路径} 映射。
    跳过外部链接（非嵌入）的图片关系，单张图片提取失败不影响整体导入。
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = Document(docx_path)
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    image_map: dict[str, str] = {}
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        if rel.is_external:
            # 外部链接图片（URL），没有内嵌数据可提取
            continue
        try:
            image_data = rel.target_part.blob
            ext = Path(rel.target_ref).suffix or ".png"
            fname = f"img_{rel.rId}{ext}"
            fpath = out / fname
            fpath.write_bytes(image_data)
            image_map[rel.rId] = str(fpath)
        except Exception:
            # 单张图片损坏/格式异常不应阻断整个文档导入
            continue

    return image_map


def _detect_block_type(text: str) -> BlockType:
    """基于内容检测 block 类型（不依赖字体信息）"""
    s = text.strip()
    if not s:
        return BlockType.PARAGRAPH

    if CHAPTER_RE.match(s):
        return BlockType.CHAPTER

    if SECTION_RE.match(s):
        return BlockType.SECTION

    if s.startswith(DIALOGUE_START) or s.endswith(DIALOGUE_END):
        return BlockType.DIALOGUE

    if re.match(r'^[\d\s]{1,6}$', s):
        return BlockType.HEADER_FOOTER

    return BlockType.PARAGRAPH


def _convert_doc_to_docx(doc_path: str) -> str:
    """
    旧版 .doc 二进制格式不能被 python-docx 打开，
    在 macOS 上用系统自带的 textutil 转换为 .docx，返回转换后的路径。
    """
    import subprocess
    out_path = str(Path(doc_path).with_suffix(".docx"))
    try:
        subprocess.run(
            ["textutil", "-convert", "docx", doc_path, "-output", out_path],
            check=True, capture_output=True, timeout=60,
        )
    except FileNotFoundError:
        raise RuntimeError(
            "检测到 .doc（旧版 Word 格式），且当前系统没有 textutil 可自动转换。\n"
            "请用 Word / Pages / LibreOffice 手动另存为 .docx 后重新导入。"
        )
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"自动转换 .doc → .docx 失败: {e.stderr.decode('utf-8', 'ignore')}\n"
            "请用 Word / Pages 手动另存为 .docx 后重新导入。"
        )
    if not Path(out_path).exists():
        raise RuntimeError("自动转换 .doc → .docx 未生成输出文件，请手动转换后重试。")
    return out_path


def import_docx(docx_path: str, verbose: bool = True) -> UnifiedDocument:
    """
    导入 DOCX 文件，转换为 UnifiedDocument。

    Args:
        docx_path: DOCX 文件路径
        verbose: 是否打印进度

    Returns:
        UnifiedDocument
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    src = Path(docx_path)
    if not src.exists():
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    # 旧版 .doc（二进制格式）自动转换
    if src.suffix.lower() == ".doc":
        if verbose:
            print(f"  🔄  检测到 .doc 旧格式，尝试自动转换为 .docx ...")
        docx_path = _convert_doc_to_docx(str(src))
        src = Path(docx_path)

    if verbose:
        print(f"📄  导入 DOCX: {docx_path}")

    try:
        docx_doc = Document(docx_path)
    except Exception as e:
        # python-docx 对非法/损坏的 zip 包会抛 PackageNotFoundError
        if type(e).__name__ == "PackageNotFoundError":
            raise ValueError(
                f"「{src.name}」不是有效的 .docx 文件（可能是旧版 .doc 被误重命名，"
                f"或文件已损坏）。请用 Word/Pages 另存为标准 .docx 后重试。"
            ) from e
        raise

    img_dir = str(Path(docx_path).parent / f"_docx_images_{Path(docx_path).stem}")
    image_map = extract_images_from_docx(docx_path, img_dir)
    if verbose and image_map:
        print(f"  🖼️  提取 {len(image_map)} 张图片")

    doc = UnifiedDocument()
    doc.metadata = Metadata(
        source_engine="docx_import",
        language="ja",
    )

    order = 0
    chapter_index = 0

    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        btype = _detect_block_type(text)

        if btype == BlockType.HEADER_FOOTER:
            continue

        block = Block(
            type=btype,
            text=text,
            ocr_raw=text,
            reading_order=order,
            confidence=0.90,
        )

        if btype == BlockType.CHAPTER:
            chapter_index += 1
            block.chapter_index = chapter_index
            from models.document import TocEntry
            doc.toc.append(TocEntry(
                title=text,
                chapter_index=chapter_index,
                block_index=len(doc.blocks),
            ))

        doc.blocks.append(block)
        order += 1

    # 处理图片引用
    for rid, img_path in image_map.items():
        last_idx = len(doc.blocks) - 1
        anchor = f"block_{last_idx}" if last_idx >= 0 else "start"
        doc.blocks.append(Block(
            type=BlockType.IMAGE_REF,
            image_path=img_path,
            image_anchor=anchor,
            reading_order=order,
        ))
        order += 1

    doc.add_log("docx_import", f"导入 {len(doc.blocks)} 个块，{chapter_index} 个章节", len(doc.blocks))

    if verbose:
        print(f"  ✅  导入完成: {len(doc.blocks)} 个块，{chapter_index} 个章节")

    return doc


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="DOCX → UnifiedDocument JSON")
    parser.add_argument("input_docx", help="输入 DOCX 文件")
    parser.add_argument("output_json", help="输出 JSON 路径")
    parser.add_argument("--quiet", "-q", action="store_true")
    args = parser.parse_args()

    doc = import_docx(args.input_docx, verbose=not args.quiet)

    with open(args.output_json, "w", encoding="utf-8") as f:
        f.write(doc.to_json())

    print(f"💾  已写入: {args.output_json}")
